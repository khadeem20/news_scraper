from selenium import webdriver
from selenium.webdriver.common.by import By #selection of elements
from selenium.webdriver.common.keys import Keys #mimic key inputs
from selenium.webdriver.chrome.service import Service #start and stop the chrome driver
from webdriver_manager.chrome import ChromeDriverManager

from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from bs4 import BeautifulSoup
from docx import Document

import time


class Sel_Scraper():

    def __init__(self,url):
        self.url = url
        self.driver = webdriver.Chrome(service=Service(ChromeDriverManager().install()))
        self.page_content= None
        self.soup= None
        self.title=None
        self.article= None
        self.doc= None

    
    def dynamic_search_url(self):
        self.driver.get(self.url)
        self.driver.implicitly_wait(10)
        # time.sleep(6) 
        self.page_content = self.driver.page_source

        self.soup = BeautifulSoup(self.page_content, 'html.parser')

        #find the title
        self.title= self.soup.find('h1', id="maincontent").get_text()


        # Find the first article with the specified class
        self.article = self.soup.find('div', class_='article__content').get_text()

        if not self.article and self.title:
            #print(self.title)
            #print(self.article)  # Print the text content of the article
       # else:
            print("Article not found.")

    def create_doc(self):
        self.doc= Document()

        self.doc.add_heading(self.title, level=1)

        self.doc.add_paragraph(self.article)

        self.doc.save('article.docx')

    def finito(self):
        self.driver.close()



    #url = 'https://edition.cnn.com/travel/article/scenic-airport-landings-2020/index.html'
    #response = requests.get(url)

    #print(url)
    # Check if the request was successful
    #if response.status_code == 200:
     #   print(response.status_code)
      #  html_content = response.text
        #html_content.find('div', class_='article__content')
       # print(html_content)



        

   


